from docx import Document
from docx.shared import Inches

document = Document()

# profile picture
document.add_picture(
    'picture.png',
    width=Inches(1.0)
    )

# name, phone number and email

name = input('what is your name?\n')
phone_number = input('what is your phone number?\n')
email = input('what is your email?\n')


document.add_paragraph(
    name + ' |' + phone_number + ' |' + email
)

# Summary

document.add_heading('SUMMARY')
document.add_paragraph(input('Tell me about yourself?\n'))

#this could also be done in a different way, by creating a variable, calling document.add_paragraph() with the variable as the parameter.
#see alternate summary below

"""document.add_heading('About me')
about_me = input('Tell me about yourself?\n')
document.add_heading(about_me)"""

# Work experience 1

document.add_heading('EXPERIENCE')
we1_paragraph = document.add_paragraph()

company_1 = input('Enter the company name\n')
position_title = input('Enter the title of the position\n')
from_date_1 = input('Enter the from date\n')
to_date_1 = input('Enter the to date\n')

we1_paragraph.add_run(company_1 + ' ').bold = True
we1_paragraph.add_run(' * ' + position_title + ' * ').bold = True
we1_paragraph.add_run(from_date_1 + '-' + to_date_1 + '\n').bold = True

experience_details1 = input('Describe your experience?\n')

we1_paragraph.add_run(experience_details1)

# Work experience 2
# see alternate work experience below

"""document.add_heading('Work experience')

company_2 = input('Enter the company name\n')
from_date_2 = input('Enter from date\n')
to_date_2 = input('Enter to date\n')
experience_details_2 = input('Describe your experience?\n')

we2_paragraph = document.add_paragraph(
    company_2 + ' ' + from_date_2 + '-' + to_date_2 + '\n' + experience_details_2)"""


# More experiences

while True:
    more_experiences = input('Do you have more experiences? Yes or No ?')
    if more_experiences.lower() == 'yes':
        we1_paragraph = document.add_paragraph()

        company_1 = input('Enter the company name\n')
        position_title = input('Enter the title of the position\n')
        from_date_1 = input('Enter the from date\n')
        to_date_1 = input('Enter the to date\n')

        we1_paragraph.add_run(company_1 + ' ').bold = True
        we1_paragraph.add_run(' * ' + position_title + ' * ').bold = True
        we1_paragraph.add_run(from_date_1 + '-' + to_date_1 + '\n').bold = True

        experience_details1 = input('Describe your experience?\n')

        we1_paragraph.add_run(experience_details1)
    else:
        break

# Education

document.add_heading('EDUCATION')
education_paragraph = document.add_paragraph()

school = input('Enter the name of the school\n')
title_of_studies = input('Enter the major\n')
from_date = input('Enter from date\n')
to_date = input('Enter to date\n')

education_paragraph.add_run(school + ' ').bold = True
education_paragraph.add_run(' * ' + title_of_studies + ' * ').bold = True
education_paragraph.add_run(from_date + '-' + to_date + '\n').bold = True

education_details = input('provide details of the education\n')

education_paragraph.add_run(education_details)

# Additional education

while True:
    additional_education = input('Do you have additional educational history? Yes or No ? \n')
    if additional_education.lower() == 'yes':
        education_paragraph = document.add_paragraph()

        school = input('Enter the name of the school\n')
        title_of_studies = input('Enter the major\n')
        from_date = input('Enter from date\n')
        to_date = input('Enter to date\n')

        education_paragraph.add_run(school + ' ').bold = True
        education_paragraph.add_run(' * ' + title_of_studies + ' * ').bold = True
        education_paragraph.add_run(from_date + '-' + to_date + '\n').bold = True

        education_details = input('provide details of the education\n')

        education_paragraph.add_run(education_details)
    else:
        break


# skills

document.add_heading('TOP SKILLS')
skill = input('Enter skill\n')
skill_paragraph = document.add_paragraph(skill)
skill_paragraph_style = 'list bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ? \n')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill\n')
        skill_paragraph = document.add_paragraph(skill)
        skill_paragraph_style = 'list bullet'
    else:
        break


document.save('cv.docx')

